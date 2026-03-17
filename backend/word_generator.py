"""Word document generation for investment memos using python-docx."""
import os
import uuid
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


def _add_styled_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(27, 79, 114)
    return heading


def _add_recommendation_badge(doc, recommendation):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"  {recommendation.upper()}  ")
    run.bold = True
    run.font.size = Pt(16)
    if recommendation.lower() == "buy":
        run.font.color.rgb = RGBColor(39, 174, 96)
    elif recommendation.lower() == "sell":
        run.font.color.rgb = RGBColor(231, 76, 60)
    else:
        run.font.color.rgb = RGBColor(243, 156, 18)
    return p


def generate_word_memo(analysis: dict, output_dir: str) -> str:
    """Generate a styled Word document investment memo.
    
    Returns the filename of the generated file.
    """
    doc = Document()
    
    # Style defaults
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(44, 62, 80)
    
    company = analysis.get("company_name", "Company")
    memo = analysis.get("memo", {}) or {}
    extracted = analysis.get("extracted_data", {}) or {}
    
    # --- Header ---
    title = doc.add_heading(f"Investment Memo — {company}", level=0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(27, 79, 114)
    
    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = subtitle_p.add_run(f"Date: {datetime.now().strftime('%B %d, %Y')}")
    date_run.font.size = Pt(11)
    date_run.font.color.rgb = RGBColor(127, 140, 141)
    
    # Recommendation badge
    recommendation = memo.get("recommendation", "Hold")
    _add_recommendation_badge(doc, recommendation)
    
    doc.add_paragraph("")  # spacer
    
    # --- Executive Summary ---
    if memo.get("executive_summary"):
        _add_styled_heading(doc, "Executive Summary", level=1)
        doc.add_paragraph(memo["executive_summary"])
    
    # --- Investment Thesis ---
    if memo.get("investment_thesis"):
        _add_styled_heading(doc, "Investment Thesis", level=1)
        doc.add_paragraph(memo["investment_thesis"])
    
    # --- Key Financial Metrics Table ---
    _add_styled_heading(doc, "Key Financial Metrics", level=1)
    
    metrics_data = []
    metric_labels = [
        ("Revenue ($M)", "revenue"),
        ("Net Income ($M)", "net_income"),
        ("EBITDA ($M)", "ebitda"),
        ("Gross Margin %", "gross_margin_pct"),
        ("Free Cash Flow ($M)", "free_cash_flow"),
    ]
    
    years_covered = analysis.get("years_covered", [])
    
    if years_covered and extracted:
        # Build table with years as columns
        num_cols = len(years_covered) + 1
        table = doc.add_table(rows=1, cols=num_cols)
        table.style = 'Light Grid Accent 1'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Header row
        hdr = table.rows[0].cells
        hdr[0].text = "Metric"
        for j, year in enumerate(years_covered):
            hdr[j + 1].text = str(year)
        
        for label, key in metric_labels:
            row = table.add_row().cells
            row[0].text = label
            items = extracted.get(key, []) or []
            year_map = {item["year"]: item["value"] for item in items if item}
            for j, year in enumerate(years_covered):
                val = year_map.get(year)
                if val is not None:
                    if "pct" in key or "margin" in key.lower():
                        row[j + 1].text = f"{val}%"
                    else:
                        row[j + 1].text = f"{val:,.0f}" if isinstance(val, (int, float)) else str(val)
                else:
                    row[j + 1].text = "N/A"
        
        # Bold header cells
        for cell in table.rows[0].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
    else:
        doc.add_paragraph("Financial metrics data not available.")
    
    # --- Valuation Summary ---
    if memo.get("valuation_summary"):
        _add_styled_heading(doc, "Valuation Summary", level=1)
        doc.add_paragraph(memo["valuation_summary"])
    
    # DCF summary if available
    dcf = analysis.get("dcf", {}) or {}
    if dcf:
        dcf_items = [
            ("Implied Share Price", f"${dcf.get('implied_price', 'N/A')}"),
            ("Enterprise Value", f"${dcf.get('enterprise_value', 'N/A')}M"),
            ("Equity Value", f"${dcf.get('equity_value', 'N/A')}M"),
            ("Upside/Downside", f"{dcf.get('upside_downside_pct', 'N/A')}%"),
        ]
        dcf_table = doc.add_table(rows=1, cols=2)
        dcf_table.style = 'Light Grid Accent 1'
        dcf_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        dcf_table.rows[0].cells[0].text = "Metric"
        dcf_table.rows[0].cells[1].text = "Value"
        for cell in dcf_table.rows[0].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        for label, value in dcf_items:
            row = dcf_table.add_row().cells
            row[0].text = label
            row[1].text = str(value)
    
    # --- Risk Factors ---
    risks = memo.get("risk_factors", [])
    if risks:
        _add_styled_heading(doc, "Risk Factors", level=1)
        for risk in risks:
            p = doc.add_paragraph(risk, style='List Bullet')
    
    # --- Footer / Disclaimer ---
    doc.add_paragraph("")  # spacer
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_p.add_run("Generated by FinAgent AI — For informational purposes only")
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(149, 165, 166)
    footer_run.italic = True
    
    # Save
    filename = f"{company.replace(' ', '_')}_memo_{uuid.uuid4().hex[:8]}.docx"
    filepath = os.path.join(output_dir, filename)
    doc.save(filepath)
    
    return filename
